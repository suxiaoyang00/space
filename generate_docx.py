import json
from pathlib import Path
from docx import Document
from docx.shared import Inches

# ========== 配置 ==========
INPUT_JSON = 'messages_with_local_paths.json'
OUTPUT_DOCX = 'report.docx'
IMAGE_DIR = 'downloaded_images'

def main():
    if not Path(INPUT_JSON).exists():
        print(f"❌ 找不到 {INPUT_JSON}，请先运行 extract_and_download_images.py")
        return

    with open(INPUT_JSON, 'r', encoding='utf-8') as f:
        messages = json.load(f)

    doc = Document()
    doc.add_heading('消息图文报告', 0)

    for msg in messages:
        msg_id = msg.get('id', '未知ID')
        createtime = msg.get('createtime', '未知时间')
        text = msg.get('text', '')

        # 添加消息头部
        doc.add_paragraph(f"ID: {msg_id} | 时间: {createtime}", style='Heading 2')

        # 添加文字内容（保留换行）
        if text:
            for line in text.split('\n'):
                doc.add_paragraph(line)

        # 插入图片
        local_paths = msg.get('local_pic_paths', [])
        for path in local_paths:
            img_path = Path(path)
            if img_path.exists():
                # 控制图片宽度（可选）
                doc.add_picture(str(img_path), width=Inches(4))
            else:
                doc.add_paragraph(f"[图片缺失: {path}]")

        # 消息之间加空行
        doc.add_paragraph()

    doc.save(OUTPUT_DOCX)
    print(f"✅ Word 报告已生成: {OUTPUT_DOCX}")
    print("💡 请用 Microsoft Word 或 WPS 打开查看。")

if __name__ == '__main__':
    main()